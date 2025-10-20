import os
import io
from typing import Dict, Any, Optional
from docx import Document
import PyPDF2
import csv
from app.models.applicant import ApplicantData, SelectionStage
from app.utils.config import settings

class FileProcessorService:
    """ファイル処理サービス（PDF, Word, CSV対応）"""
    
    def __init__(self):
        pass
    
    async def process_file(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """
        ファイルを処理してテキストを抽出
        
        Args:
            file_content: ファイルの内容（バイト）
            filename: ファイル名
            file_type: ファイルタイプ（application/pdf, application/vnd.openxmlformats-officedocument.wordprocessingml.document, text/csv）
        
        Returns:
            抽出結果
        """
        try:
            if file_type == "application/pdf" or filename.endswith('.pdf'):
                return await self._process_pdf(file_content)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.endswith('.docx'):
                return await self._process_word(file_content)
            elif file_type == "text/csv" or filename.endswith('.csv'):
                return await self._process_csv(file_content)
            else:
                return {
                    "success": False,
                    "error": f"サポートされていないファイル形式です: {file_type}"
                }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _process_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """PDFファイルを処理"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            extracted_text = ""
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n\n"
            
            return {
                "success": True,
                "text": extracted_text,
                "page_count": len(pdf_reader.pages),
                "confidence": 0.9
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF処理エラー: {str(e)}"
            }
    
    async def _process_word(self, file_content: bytes) -> Dict[str, Any]:
        """Wordファイルを処理"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            extracted_text = ""
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            # テーブルも処理
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + "\t"
                    extracted_text += "\n"
            
            return {
                "success": True,
                "text": extracted_text,
                "page_count": len(doc.paragraphs),
                "confidence": 0.95
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Word処理エラー: {str(e)}"
            }
    
    async def _process_csv(self, file_content: bytes) -> Dict[str, Any]:
        """CSVファイルを処理"""
        try:
            csv_text = file_content.decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(csv_text))
            
            rows = list(csv_reader)
            
            # CSVの場合は行ごとに処理
            return {
                "success": True,
                "csv_data": rows,
                "row_count": len(rows),
                "confidence": 1.0
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"CSV処理エラー: {str(e)}"
            }

class AICategorizationService:
    """AI自動仕分けサービス"""
    
    def __init__(self):
        pass
    
    async def categorize_and_format(self, extracted_text: str) -> Dict[str, Any]:
        """
        抽出されたテキストをAIで仕分け・整形
        
        Args:
            extracted_text: 抽出されたテキスト
        
        Returns:
            仕分け結果と整形されたデータ
        """
        from vertexai.generative_models import GenerativeModel
        import vertexai
        import json
        import os
        
        try:
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = settings.google_application_credentials
            vertexai.init(project=settings.google_cloud_project_id, location="us-central1")
            
            model = GenerativeModel("gemini-1.5-flash")
            
            prompt = f"""
以下は応募書類から抽出されたテキストです。このテキストを分析し、以下の作業を行ってください：

1. 応募者データの構造化（JSON形式）
2. 適切な選考ステージの推奨
3. 不足している情報のリスト化

【抽出テキスト】
{extracted_text}

【出力形式】
以下のJSON形式で出力してください：

{{
  "applicant_data": {{
    "name": "氏名",
    "email": "メールアドレス",
    "phone": "電話番号",
    "education": [{{"institution": "学校名", "degree": "学位", "field": "専攻", "year": "卒業年"}}],
    "work_experience": [{{"company": "会社名", "position": "役職", "duration": "期間", "description": "業務内容"}}],
    "technical_skills": ["スキル1", "スキル2"],
    "soft_skills": ["スキル1", "スキル2"],
    "certifications": ["資格1", "資格2"],
    "motivation": "志望動機",
    "career_goals": "キャリア目標",
    "additional_info": "その他の情報"
  }},
  "recommended_stage": "document_screening / first_interview / second_interview",
  "recommendation_reason": "推奨理由（1-2文）",
  "missing_info": ["不足している情報1", "不足している情報2"],
  "quality_score": 8.5,
  "auto_actions": [
    {{"action": "アクション名", "description": "説明"}}
  ]
}}

【選考ステージの選び方】
- document_screening: 書類のみで評価が必要
- first_interview: 書類は良好、一次面接へ進める
- second_interview: 優秀な候補者、二次面接から開始
- rejected: 明らかに基準に達していない
"""

            response = model.generate_content(prompt)
            result_text = response.text
            
            # JSONパース
            if "```json" in result_text:
                json_start = result_text.find("```json") + 7
                json_end = result_text.find("```", json_start)
                result_text = result_text[json_start:json_end].strip()
            elif "```" in result_text:
                json_start = result_text.find("```") + 3
                json_end = result_text.find("```", json_start)
                result_text = result_text[json_start:json_end].strip()
            
            result = json.loads(result_text)
            
            return {
                "success": True,
                "data": result
            }
            
        except Exception as e:
            print(f"AI仕分けエラー: {str(e)}")
            # エラー時はデフォルト値を返す
            return {
                "success": False,
                "error": str(e),
                "data": {
                    "applicant_data": {},
                    "recommended_stage": "document_screening",
                    "recommendation_reason": "自動分析に失敗しました。手動で確認してください。",
                    "missing_info": ["すべての情報を手動で入力してください"],
                    "quality_score": 0,
                    "auto_actions": []
                }
            }
